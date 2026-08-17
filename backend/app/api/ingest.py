"""
api/ingest.py — Fast, reliable multi-media file ingestion.

Design:
  - NEVER calls an LLM during ingestion (no Groq/Gemini rate limits)
  - Stores chunks DIRECTLY into MongoDB (mem0_memories collection)
  - Parser runs locally — no external API calls for text extraction
  - Gemini Vision for images is OPTIONAL — skipped gracefully if unavailable

Supported:
  PDF   → LlamaParse (cloud, layout-aware) with pypdf fallback
  DOCX  → python-docx (local, tables included)
  XLSX  → openpyxl (local, all sheets)
  TXT   → direct read
  PNG/JPG/WEBP/GIF → Gemini Vision description OR filename+size metadata
  MP3/WAV/M4A → metadata stored (audio transcription future feature)
  MP4/MOV → metadata stored (video future feature)
"""
from __future__ import annotations

import io
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import List

from fastapi import APIRouter, HTTPException, UploadFile, File, Depends
from pydantic import BaseModel

from app.auth import get_current_user_id
from app.config import GEMINI_API_KEY, GEMINI_MODEL

router = APIRouter()

MAX_FILE_SIZE = 50 * 1024 * 1024   # 50 MB
CHUNK_SIZE    = 1500
CHUNK_OVERLAP = 150

# ── Supported types ───────────────────────────────────────────────────────────
EXT_MAP = {
    ".pdf":  "pdf",
    ".docx": "docx", ".doc": "docx",
    ".xlsx": "xlsx", ".xls": "xlsx",
    ".txt":  "txt",  ".md": "txt", ".csv": "txt",
    ".png":  "image", ".jpg": "image", ".jpeg": "image",
    ".webp": "image", ".gif": "image", ".bmp": "image",
    ".mp3":  "audio", ".wav": "audio", ".m4a": "audio", ".ogg": "audio",
    ".mp4":  "video", ".mov": "video", ".avi": "video",
}

MIME_MAP = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xlsx",
    "text/plain": "txt", "text/markdown": "txt", "text/csv": "txt",
    "image/png": "image", "image/jpeg": "image",
    "image/webp": "image", "image/gif": "image",
    "audio/mpeg": "audio", "audio/wav": "audio", "audio/mp4": "audio",
    "video/mp4": "video", "video/quicktime": "video",
}


# ── Response ──────────────────────────────────────────────────────────────────
class IngestResponse(BaseModel):
    filename:       str
    file_type:      str
    chunks_stored:  int
    message:        str


# ── Chunker ───────────────────────────────────────────────────────────────────
def _chunk(text: str, source: str) -> list[str]:
    text = text.strip()
    if not text:
        return []
    out, start = [], 0
    while start < len(text):
        out.append(f"[Source: {source}]\n{text[start:start + CHUNK_SIZE]}")
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return out


# ── Direct MongoDB store (NO LLM) ─────────────────────────────────────────────
async def _store_chunks(chunks: list[str], user_id: str, filename: str, file_type: str) -> int:
    """Insert chunks directly into mem0_memories — no LLM extraction, instant."""
    from app.db.mongo import get_db
    db = get_db()
    if db is None:
        raise RuntimeError("MongoDB not connected")
    docs = [
        {
            "memory":     chunk,
            "user_id":    user_id,
            "hash":       uuid.uuid4().hex,
            "created_at": datetime.now(timezone.utc),
            "metadata":   {"source": filename, "file_type": file_type},
        }
        for chunk in chunks
    ]
    result = await db["mem0_memories"].insert_many(docs)
    return len(result.inserted_ids)


# ── Parsers ───────────────────────────────────────────────────────────────────

def _parse_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _parse_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = []
    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    # Tables → markdown
    for table in doc.tables:
        rows = []
        for i, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
            if i == 0:
                rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def _parse_xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows, headers = [], []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            vals = [str(v).strip() if v is not None else "" for v in row]
            if not any(vals):
                continue
            if i == 0:
                headers = vals
                rows.append("| " + " | ".join(vals) + " |")
                rows.append("| " + " | ".join(["---"] * len(vals)) + " |")
            else:
                if headers:
                    pairs = [f"{h}={v}" for h, v in zip(headers, vals) if v]
                    rows.append(", ".join(pairs))
                else:
                    rows.append(" | ".join(vals))
        if rows:
            parts.append(f"[Sheet: {name}]\n" + "\n".join(rows))
    return "\n\n".join(parts)


def _parse_pdf_local(data: bytes) -> str:
    """Fast local PDF text extraction via pypdf (no API call)."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {i+1}]\n{text.strip()}")
    return "\n\n".join(pages)


def _parse_pdf_llamaparse(data: bytes, filename: str) -> str:
    """Layout-aware PDF via LlamaParse — used only if key is available."""
    from app.config import LLAMA_CLOUD_API_KEY
    if not LLAMA_CLOUD_API_KEY:
        return _parse_pdf_local(data)
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        from app.documents.parser import parse_pdf as _parse
        elements = _parse(tmp_path)
        if not elements:
            return _parse_pdf_local(data)
        parts = []
        for e in elements:
            if e["type"] in ("heading", "paragraph"):
                parts.append(e["content"])
            elif e["type"] == "table":
                c = e["content"]
                parts.append(c if isinstance(c, str) else str(c))
            # skip images in flat ingest — they're handled by documents/upload
        return "\n\n".join(p for p in parts if p.strip())
    except Exception:
        return _parse_pdf_local(data)
    finally:
        Path(tmp_path).unlink(missing_ok=True)


async def _parse_image(data: bytes, filename: str) -> str:
    """Describe image via Gemini Vision. Tries each model with individual timeouts."""
    if not GEMINI_API_KEY:
        return f"[Image: {filename} — {len(data)//1024}KB — vision description unavailable, Gemini key not set]"
    import base64
    import httpx
    ext  = Path(filename).suffix.lower()
    mime = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".webp": "image/webp", ".gif": "image/gif"}.get(ext, "image/jpeg")
    b64  = base64.b64encode(data).decode()
    payload = {"contents": [{"parts": [
        {"inline_data": {"mime_type": mime, "data": b64}},
        {"text": "Describe this image factually. If it contains text, numbers, charts, or diagrams, extract all of them verbatim."},
    ]}]}
    models_to_try = ["gemini-3.6-flash", "gemini-2.5-flash-preview-05-20", "gemini-1.5-flash"]
    # Each model gets its own 30s timeout — if one times out, move to next immediately
    for model in models_to_try:
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GEMINI_API_KEY}"
                r = await client.post(url, json=payload)
                if r.status_code == 200:
                    text_out = r.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
                    if text_out:
                        return f"[Image Content ({filename})]: {text_out}"
                elif r.status_code in (429, 503):
                    continue  # rate limit or overload — try next model
        except (httpx.TimeoutException, httpx.ConnectError):
            continue  # this model timed out — try next immediately
        except Exception:
            continue
    size = f"{len(data)//1024}KB"
    return f"[Image: {filename} ({size}) — vision description unavailable]"


def _media_metadata(filename: str, data: bytes, file_type: str) -> str:
    """Return metadata string for audio/video (full transcription is future work)."""
    size = f"{len(data)/(1024*1024):.1f}MB"
    return (
        f"[{file_type.upper()} file uploaded: {filename} ({size})]\n"
        f"Note: Automatic transcription coming soon. "
        f"You can ask the assistant about this file by name."
    )


# ── Main endpoint ─────────────────────────────────────────────────────────────
@router.post("/ingest", response_model=IngestResponse)
async def ingest_file(
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user_id),
):
    data     = await file.read()
    filename = file.filename or "upload"
    ext      = Path(filename).suffix.lower()

    if len(data) > MAX_FILE_SIZE:
        raise HTTPException(400, f"File exceeds {MAX_FILE_SIZE//(1024*1024)}MB limit.")

    # Detect type
    file_type = MIME_MAP.get(file.content_type or "") or EXT_MAP.get(ext)
    if not file_type:
        raise HTTPException(400, f"Unsupported file type: {ext or file.content_type}")

    # Parse → text
    try:
        if file_type == "txt":
            text = _parse_txt(data)
        elif file_type == "docx":
            text = _parse_docx(data)
        elif file_type == "xlsx":
            text = _parse_xlsx(data)
        elif file_type == "pdf":
            # For inline chat attach: use fast local pypdf (no API, instant)
            # Use /api/documents/upload for full LlamaParse pipeline
            text = _parse_pdf_local(data)
        elif file_type == "image":
            text = await _parse_image(data, filename)
        elif file_type in ("audio", "video"):
            text = _media_metadata(filename, data, file_type)
        else:
            raise HTTPException(400, f"Unsupported: {file_type}")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Parse error for {filename}: {e}")

    if not text.strip():
        raise HTTPException(422, f"No content extracted from {filename}.")

    # Store directly — NO LLM, instant
    chunks  = _chunk(text, filename)
    try:
        stored = await _store_chunks(chunks, user_id, filename, file_type)
    except Exception as e:
        raise HTTPException(500, f"Storage failed: {e}")

    return IngestResponse(
        filename=filename,
        file_type=file_type,
        chunks_stored=stored,
        message=f"✅ '{filename}' ingested — {stored} chunk(s) stored. The assistant can now recall this content.",
    )


@router.post("/ingest/multiple", response_model=List[IngestResponse])
async def ingest_multiple(
    files: List[UploadFile] = File(...),
    user_id: str = Depends(get_current_user_id),
):
    results = []
    for f in files:
        try:
            results.append(await ingest_file(file=f, user_id=user_id))
        except HTTPException as e:
            results.append(IngestResponse(
                filename=f.filename or "unknown",
                file_type="unknown",
                chunks_stored=0,
                message=f"Error: {e.detail}",
            ))
    return results
